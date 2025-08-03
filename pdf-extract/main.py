import fitz  # PyMuPDF
import pymupdf  # PyMuPDF
import pdfplumber
import pandas as pd
import os
from docx import Document
import email
import zipfile
from pathlib import Path
import win32com.client

def create_output_structure(file_path):
    """Create folder structure based on the input file name"""
    base_name = Path(file_path).stem  # Get filename without extension
    base_folder = f"extracted_{base_name}"
    
    # Create main folder and subfolders
    folders = {
        'main': base_folder,
        'text': os.path.join(base_folder, 'text'),
        'tables': os.path.join(base_folder, 'tables'),
        'images': os.path.join(base_folder, 'images')
    }
    
    for folder in folders.values():
        os.makedirs(folder, exist_ok=True)
    
    return folders

def extract_from_pdf(pdf_path, folders):
    """Extract text, tables, and images from PDF"""
    print(f"Processing PDF: {pdf_path}")
    
    # === TEXT EXTRACTION ===
    print("Extracting text from PDF...")
    doc = fitz.open(pdf_path)
    text_file = os.path.join(folders['text'], "pdf_text.txt")
    
    with open(text_file, "wb") as out:
        for page in doc:
            text = page.get_text().encode("utf8")
            out.write(text)
            out.write(bytes((12,)))  # page delimiter
    print("PDF text extraction completed!")

    # === IMAGE EXTRACTION ===
    print("Extracting images from PDF...")
    image_count = 0
    for page_index in range(len(doc)):
        page = doc[page_index]
        image_list = page.get_images()

        if image_list:
            print(f"Found {len(image_list)} images on page {page_index + 1}")
        
        for image_index, img in enumerate(image_list, start=1):
            xref = img[0]
            pix = pymupdf.Pixmap(doc, xref)

            if pix.n - pix.alpha > 3:
                pix = pymupdf.Pixmap(pymupdf.csRGB, pix)

            image_file = os.path.join(folders['images'], f"pdf_page_{page_index + 1}_image_{image_index}.png")
            pix.save(image_file)
            pix = None
            image_count += 1
    
    print(f"PDF image extraction completed! ({image_count} images)")

    # === TABLE EXTRACTION ===
    print("Extracting tables from PDF...")
    with pdfplumber.open(pdf_path) as pdf:
        all_tables_data = []
        
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            
            if tables:
                print(f"Found {len(tables)} tables on page {page_num + 1}")
                
                for table_num, table in enumerate(tables):
                    if table and len(table) > 0:
                        # Create DataFrame
                        if len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                        else:
                            df = pd.DataFrame(table)
                        
                        df = df.fillna("")
                        
                        # Save table files
                        table_csv = os.path.join(folders['tables'], f"pdf_page_{page_num + 1}_table_{table_num + 1}.csv")
                        table_excel = os.path.join(folders['tables'], f"pdf_page_{page_num + 1}_table_{table_num + 1}.xlsx")
                        
                        df.to_csv(table_csv, index=False, encoding='utf-8')
                        df.to_excel(table_excel, index=False)
                        
                        table_info = {
                            'source': 'PDF',
                            'page': page_num + 1,
                            'table_num': table_num + 1,
                            'rows': len(df),
                            'columns': len(df.columns),
                            'filename': table_csv
                        }
                        all_tables_data.append(table_info)
        
        # Save table summary
        if all_tables_data:
            summary_file = os.path.join(folders['tables'], "pdf_tables_summary.csv")
            summary_df = pd.DataFrame(all_tables_data)
            summary_df.to_csv(summary_file, index=False)
            print(f"PDF tables extracted: {len(all_tables_data)}")
    
    doc.close()

def extract_from_docx(docx_path, folders):
    """Extract text, tables, and images from DOCX"""
    print(f"Processing DOCX: {docx_path}")
    
    doc = Document(docx_path)
    
    # === TEXT EXTRACTION ===
    print("Extracting text from DOCX...")
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    
    text_file = os.path.join(folders['text'], "docx_text.txt")
    with open(text_file, "w", encoding="utf-8") as f:
        f.write("\n".join(text_content))
    print("DOCX text extraction completed!")
    
    # === TABLE EXTRACTION ===
    print("Extracting tables from DOCX...")
    table_count = 0
    all_tables_data = []
    
    for table_index, table in enumerate(doc.tables):
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        if data:
            if len(data) > 1:
                df = pd.DataFrame(data[1:], columns=data[0])
            else:
                df = pd.DataFrame(data)
            
            table_csv = os.path.join(folders['tables'], f"docx_table_{table_index + 1}.csv")
            table_excel = os.path.join(folders['tables'], f"docx_table_{table_index + 1}.xlsx")
            
            df.to_csv(table_csv, index=False, encoding='utf-8')
            df.to_excel(table_excel, index=False)
            
            table_info = {
                'source': 'DOCX',
                'page': 'N/A',
                'table_num': table_index + 1,
                'rows': len(df),
                'columns': len(df.columns),
                'filename': table_csv
            }
            all_tables_data.append(table_info)
            table_count += 1
    
    if all_tables_data:
        summary_file = os.path.join(folders['tables'], "docx_tables_summary.csv")
        summary_df = pd.DataFrame(all_tables_data)
        summary_df.to_csv(summary_file, index=False)
    
    print(f"DOCX tables extracted: {table_count}")
    
    # === IMAGE EXTRACTION ===
    print("Extracting images from DOCX...")
    image_count = 0
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
            for i, media_file in enumerate(media_files):
                with docx_zip.open(media_file) as img_file:
                    img_data = img_file.read()
                    file_ext = media_file.split('.')[-1]
                    image_file = os.path.join(folders['images'], f"docx_image_{i + 1}.{file_ext}")
                    with open(image_file, "wb") as f:
                        f.write(img_data)
                    image_count += 1
    except Exception as e:
        print(f"Could not extract images from DOCX: {e}")
    
    print(f"DOCX image extraction completed! ({image_count} images)")

def extract_from_eml(eml_path, folders):
    """Extract content from EML email file"""
    print(f"Processing EML: {eml_path}")
    
    with open(eml_path, 'rb') as f:
        msg = email.message_from_bytes(f.read())
    
    # Extract email metadata and content
    subject = msg.get('Subject', 'No Subject')
    sender = msg.get('From', 'Unknown Sender')
    date = msg.get('Date', 'Unknown Date')
    
    # Extract body
    body = ""
    attachment_count = 0
    
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get("Content-Disposition"))
            
            if content_type == "text/plain" and "attachment" not in content_disposition:
                body += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            elif "attachment" in content_disposition:
                filename = part.get_filename()
                if filename:
                    attachment_file = os.path.join(folders['images'], f"email_attachment_{attachment_count + 1}_{filename}")
                    with open(attachment_file, "wb") as f:
                        f.write(part.get_payload(decode=True))
                    attachment_count += 1
    else:
        body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
    
    # Save email content
    text_file = os.path.join(folders['text'], "email_content.txt")
    with open(text_file, "w", encoding="utf-8") as f:
        f.write(f"Subject: {subject}\n")
        f.write(f"From: {sender}\n")
        f.write(f"Date: {date}\n")
        f.write(f"Body:\n{body}")
    
    print(f"Email extraction completed! ({attachment_count} attachments)")

def extract_from_msg(msg_path, folders):
    """Extract content from MSG email file"""
    print(f"Processing MSG: {msg_path}")
    
    try:
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        msg = outlook.OpenSharedItem(msg_path)
        
        # Extract content
        subject = msg.Subject or "No Subject"
        sender = msg.SenderName or "Unknown Sender"
        body = msg.Body or ""
        
        # Save content
        text_file = os.path.join(folders['text'], "msg_content.txt")
        with open(text_file, "w", encoding="utf-8") as f:
            f.write(f"Subject: {subject}\n")
            f.write(f"From: {sender}\n")
            f.write(f"Body:\n{body}")
        
        print("MSG extraction completed!")
        
    except Exception as e:
        print(f"Error processing MSG file: {e}")

def process_document(file_path):
    """Main function to process any supported document type"""
    file_ext = Path(file_path).suffix.lower()
    
    if not os.path.exists(file_path):
        print(f"Error: File not found - {file_path}")
        return
    
    # Create output folder structure
    folders = create_output_structure(file_path)
    print(f"Created output folder: {folders['main']}")
    
    # Process based on file type
    if file_ext == '.pdf':
        extract_from_pdf(file_path, folders)
    elif file_ext == '.docx':
        extract_from_docx(file_path, folders)
    elif file_ext == '.eml':
        extract_from_eml(file_path, folders)
    elif file_ext == '.msg':
        extract_from_msg(file_path, folders)
    else:
        print(f"Unsupported file type: {file_ext}")
        print("Supported formats: .pdf, .docx, .eml, .msg")
        return
    
    print(f"\nExtraction completed! Check the folder: {folders['main']}")
    print(f"├── text/     - Text content")
    print(f"├── tables/   - Extracted tables (CSV & Excel)")
    print(f"└── images/   - Images and attachments")

# === MAIN EXECUTION ===
if __name__ == "__main__":
    # Example usage - change this to your file path
    file_path = "pdf-extract\\game-of-thrones.pdf"

    # You can also process multiple files
    files_to_process = [
        "pdf-extract\\game-of-thrones.pdf",
        # "path\\to\\your\\document.docx",
        # "path\\to\\your\\email.eml",
        # "path\\to\\your\\email.msg"
    ]
    
    for file_path in files_to_process:
        if os.path.exists(file_path):
            process_document(file_path)
            print("\n" + "="*50 + "\n")
        else:
            print(f"File not found: {file_path}")